#!/usr/bin/env python3
# -*- coding: utf-8 -*-
r"""
md2docx_mermaid.py  (v2.3 — sane unicode & entities)

Uso:
  # 1 fichero .md
  python md2docx.py input.md salida.docx

  # carpeta con varios .md numerados 01_xxx.md, 02_xxx.md, ...
  python md2docx.py carpeta_md salida.docx
"""
import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
from pathlib import Path

import pypandoc
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------------------------------------------------
# regex básicos
# ----------------------------------------------------------------------
MERMAID_BLOCK_RE = re.compile(r"```mermaid\s*(.*?)\s*```", re.DOTALL)

# Soporte carpeta con múltiples .md numerados 01_xxx.md
MD_ORDER_RE = re.compile(r"^(\d+)_.*\.md$", re.IGNORECASE)


def collect_md_files_from_dir(dir_path: Path):
    if not dir_path.is_dir():
        raise ValueError(f"{dir_path} no es una carpeta válida")

    candidates = []
    for p in dir_path.iterdir():
        if not p.is_file():
            continue
        m = MD_ORDER_RE.match(p.name)
        if m:
            num = int(m.group(1))
            candidates.append((num, p.name, p))

    candidates.sort(key=lambda t: (t[0], t[1]))
    return [c[2] for c in candidates]


def build_combined_markdown_from_dir(dir_path: Path) -> str:
    files = collect_md_files_from_dir(dir_path)
    if not files:
        raise RuntimeError(
            f"No se han encontrado ficheros .md válidos en {dir_path} "
            "(se buscan nombres tipo '01_xxx.md', '02_algo.md', ...)."
        )

    parts = []
    for f in files:
        text = f.read_text(encoding="utf-8")
        text = clean_markdown_text(text)
        text = fix_acceptance_criteria_formatting(text)
        parts.append(text.strip() + "\n")

    return "\n\n".join(parts).rstrip() + "\n"


# ----------------------------------------------------------------------
# helpers mmdc
# ----------------------------------------------------------------------
def which_any(cands):
    for c in cands:
        p = shutil.which(c)
        if p:
            return p
    return None


def all_mmdc_candidates():
    """Genera candidatos de comando para mermaid-cli."""
    # 1) mmdc directo
    for base in ["mmdc", "mmdc.cmd", "mmdc.exe"]:
        p = which_any([base])
        if p:
            yield [p]

    # 2) node_modules locales
    local_paths = [
        Path("node_modules") / ".bin" / "mmdc",
        Path("node_modules") / ".bin" / "mmdc.cmd",
        Path("node_modules") / ".bin" / "mmdc.exe",
    ]
    for lp in local_paths:
        if lp.exists() and os.access(lp, os.X_OK):
            yield [str(lp)]

    # 3) npx (diferentes variantes)
    npx_bin = which_any(["npx", "npx.cmd", "npx.exe"])
    if npx_bin:
        yield [npx_bin, "-y", "-p", "@mermaid-js/mermaid-cli", "mmdc"]
        yield [npx_bin, "-y", "@mermaid-js/mermaid-cli", "mmdc"]
        yield [npx_bin, "-y", "@mermaid-js/mermaid-cli"]


def try_run(cmd, args):
    env = os.environ.copy()
    return subprocess.run(
        cmd + args,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        shell=False,
        env=env,
    )


# ---------- Normalizadores base ----------
_SUBGRAPH_LINE = re.compile(
    r"^(?P<indent>\s*)subgraph\s+(?P<body>.+?)\s*$", re.MULTILINE
)


def _strip_accents(s: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c)
    )


def _sanitize_id(text: str) -> str:
    text = _strip_accents(text)
    text = re.sub(r"\s+", "_", text.strip())
    text = re.sub(r"[^A-Za-z0-9_]", "", text)
    if not text:
        text = "sg"
    if text[0].isdigit():
        text = "sg_" + text
    return text


def fix_mermaid_subgraphs(mmd: str) -> str:
    def repl(match):
        indent = match.group("indent")
        body = match.group("body").strip()
        if re.match(r"^\w+\s*\[.*\]\s*$", body):
            return match.group(0)
        raw = body
        body = re.sub(r'^\s*["\'](.+?)["\']\s*$', r"\1", body)
        sg_id = _sanitize_id(body)
        label = raw.strip()
        return f'{indent}subgraph {sg_id}["{label}"]'

    return _SUBGRAPH_LINE.sub(repl, mmd)


# ---------- Normalización UNICODE y limpieza de caracteres raros ----------
_CONTROL_CHARS = (
    "".join(chr(c) for c in range(0, 32) if c not in (9, 10, 13)) + chr(127)
)
_CONTROL_RE = re.compile(f"[{re.escape(_CONTROL_CHARS)}]")


def normalize_unicode(text: str) -> str:
    text = unicodedata.normalize("NFKC", text)
    text = _CONTROL_RE.sub("", text)
    return text

def clean_markdown_text(text: str) -> str:
    # Eliminar BOM en cualquier posición
    text = text.replace("\ufeff", "")
    return text


def fix_acceptance_criteria_formatting(text: str) -> str:
    """
    Formatea y limpia el markdown antes de la conversión.
    1. Rompe estructuras de tabla incorrectas que causan texto vertical en Word.
    2. Separa listas concatenadas (HU-045/077).
    3. Asegura el espaciado correcto para los Criterios de Aceptación.
    4. Convierte <br> a marcador para post-procesamiento.
    """
    
    # --- FASE 0: Convertir <br> a carácter especial ---
    # Usamos el carácter LINE SEPARATOR (U+2028) que Pandoc preserva como texto
    # y luego lo convertimos a salto de línea real en Word
    text = re.sub(r'<br\s*/?>', '\u2028', text)
    
    # --- FASE 1: Limpieza General de Estructura ---
    
    # Eliminar líneas de metadatos del origen que confunden a Pandoc
    text = re.sub(r'^The following table:.*$', '', text, flags=re.MULTILINE)

    # FIX CRÍTICO: Asegurar que los encabezados (###) tengan saltos de línea antes
    # Esto evita que un título pegado a texto anterior se rompa.
    text = re.sub(r'([^\n])\s*(#{1,6}\s)', r'\1\n\n\2', text)

    # FIX CRÍTICO: Desconcatenar ítems de lista pegados (Problema HU-045/077)
    # Convierte "Texto - [ ] Item" en "Texto\n- [ ] Item"
    text = re.sub(r'([^\n])\s+-\s+\[([ xX]?)\]', r'\1\n- [\2]', text)

    lines = text.split('\n')
    fixed_lines = []
    
    for line in lines:
        stripped = line.strip()
        
        # --- FASE 2: Corrección de "Texto Vertical" (Tablas rotas) ---
        # Si una línea parece una fila de tabla (| ### ...) pero contiene un Título,
        # eliminamos las tuberías (|) para convertirlo en texto normal.
        if stripped.startswith('|') and '#' in stripped:
            line = line.replace('|', '')
            stripped = line.strip()

        # --- FASE 3: Limpieza de comillas en listas ---
        # El origen a veces trae comillas envolviendo la lista: "- [ ] Texto"
        if stripped.startswith('"- [') or stripped.startswith("'- ["):
            line = line.replace('"- [', '- [').replace("'- [", '- [')
        if stripped.endswith('"') and '- [' in line:
            line = line.rstrip('"')

        # --- FASE 4: Formateo de Criterios de Aceptación ---
        if ('Criterios de aceptación:' in line or 'Criterios de aceptación:**' in line):
            # Caso A: Cabecera e ítem en la misma línea -> "Criterios: - [ ] Item"
            if ' - [' in line:
                parts = line.split(' - [', 1)
                header = parts[0].strip()
                content = '- [' + parts[1]
                fixed_lines.append(header)
                fixed_lines.append('') # Línea vacía OBLIGATORIA para que Pandoc haga lista
                fixed_lines.append(content)
            # Caso B: Cabecera normal
            else:
                fixed_lines.append(line)
                fixed_lines.append('') # Forzamos línea vacía por seguridad
        else:
            fixed_lines.append(line)
    
    return '\n'.join(fixed_lines)

# ---------- Fix escapes y entidades dentro de nodos ----------
_NODE_LABEL_RE = re.compile(
    r"(\w+\s*(?:\([^\)]*\))?\s*(?:\[[^\]]*\]|\([^\)]*\)|\{[^\}]*\}|\"[^\"]*\"|'[^']*'))"
)


def _escape_xml_entities_in_label(label: str) -> str:
    def repl(m):
        inner = m.group(0)
        inner = inner.replace("&", "&amp;")
        inner = inner.replace("<", "&lt;")
        inner = inner.replace(">", "&gt;")
        return inner

    return re.sub(r"(?<!&)(?:&(?!(?:amp|lt|gt);))|[<>]", repl, label)


def escape_entities_in_nodes(mmd: str) -> str:
    def repl(match):
        full = match.group(0)
        return _escape_xml_entities_in_label(full)

    return _NODE_LABEL_RE.sub(repl, mmd)


# ---------- Fix de líneas dobles y paréntesis en Mermaid ----------
_LINE_CONT_RE = re.compile(r"\\\n")


def fix_double_lines(mmd: str) -> str:
    return _LINE_CONT_RE.sub(" ", mmd)


_PARENS_LABEL_RE = re.compile(r"(\[[^\]]*\]|\([^\)]*\))")


def escape_parens_in_labels(mmd: str) -> str:
    """
    Escapa paréntesis SOLO dentro de etiquetas de nodo estándar.
    NO escapa:
    - Nodos terminales de flowchart: ([texto]) o (([texto]))
    - Flechas con etiquetas: -->|texto|
    - Subgraphs
    """
    # DESHABILITADO: Esta función causaba problemas con sintaxis flowchart moderna
    # Los paréntesis en Mermaid moderno no necesitan escape en la mayoría de casos
    return mmd


# ---------- Fix edge labels ----------
_EDGE_LABEL_RE = re.compile(r"(--+)([^-]*?)(--+>?)")


def fix_edge_labels(mmd: str) -> str:
    """
    Añade etiquetas tipo |label| a aristas del estilo:
        A --> texto B
    pero NO toca:
        - Aristas que ya tienen |label|
        - Aristas con [atributos]
    """
    def repl(m):
        left, label, right = m.group(1), m.group(2), m.group(3)

        if "|" in label:
            return m.group(0)
        if "[" in label or "]" in label:
            return m.group(0)

        text = label.strip()
        if not text:
            return m.group(0)

        return f"{left}|{text}|{right}"

    return _EDGE_LABEL_RE.sub(repl, mmd)


# ---------- Mermaid rendering (devuelve True/False, nunca lanza) ----------
def fix_parens_in_rectangular_nodes(mmd: str) -> str:
    """
    Escapa paréntesis dentro de nodos rectangulares [texto] para evitar errores de parsing.
    Convierte: A[Seleccionar archivo(s)] -> A["Seleccionar archivo(s)"]
    Solo afecta a nodos que contienen paréntesis y no están ya entrecomillados.
    """
    # Regex para detectar nodos rectangulares: ID[contenido]
    # Capturamos: ID, apertura [, contenido, cierre ]
    node_pattern = re.compile(r'(\w+)\[([^\]]+)\]')
    
    def repl(match):
        node_id = match.group(1)
        content = match.group(2).strip()
        
        # Si ya tiene comillas dobles o simples al inicio, no tocar
        if content.startswith('"') or content.startswith("'"):
            return match.group(0)
        
        # Si contiene paréntesis, envolver en comillas dobles
        if '(' in content or ')' in content:
            return f'{node_id}["{content}"]'
        
        return match.group(0)
    
    return node_pattern.sub(repl, mmd)


def render_mermaid_to_png(
    code: str, out_png: Path, width: int, background: str, theme: str
) -> bool:
    """
    Devuelve True si la imagen se renderiza correctamente.
    Devuelve False si mmdc falla (Invalid expression, etc.).
    """
    code = normalize_unicode(code)
    # Limpiar tags HTML que causan problemas en Mermaid
    code = re.sub(r'<br\s*/?>', ' ', code)  # Convertir <br/> a espacio simple
    # Normalizar espacios múltiples a un solo espacio (excepto indentación)
    lines = code.split('\n')
    normalized_lines = []
    for line in lines:
        # Preservar indentación pero limpiar espacios extra en el contenido
        stripped = line.lstrip()
        indent = line[:len(line) - len(stripped)]
        # Reemplazar múltiples espacios por uno solo en el contenido
        content = re.sub(r' {2,}', ' ', stripped)
        normalized_lines.append(indent + content)
    code = '\n'.join(normalized_lines)
    code = fix_mermaid_subgraphs(code)
    code = fix_parens_in_rectangular_nodes(code)  # Arreglar paréntesis en nodos rectangulares
    code = fix_double_lines(code)
    code = escape_parens_in_labels(code)
    # code = escape_entities_in_nodes(code)  # DESHABILITADO: rompe sintaxis -->
    # code = fix_edge_labels(code)  # DESHABILITADO: rompe sintaxis -->|"label"|

    with tempfile.NamedTemporaryFile(
        delete=False, suffix=".mmd", mode="w", encoding="utf-8"
    ) as tmp:
        tmp.write(code)
        tmp_path = tmp.name

    last_err = None
    ok = False

    for base_cmd in all_mmdc_candidates():
        cmd = base_cmd + [
            "-i",
            tmp_path,
            "-o",
            str(out_png),
            "-w",
            str(width),
            "-s",
            "2",
            "-b",
            background,
            "-t",
            theme,
        ]
        try:
            res = try_run(base_cmd, cmd[len(base_cmd) :])
            if res.returncode == 0:
                ok = True
                break
            else:
                last_err = res.stderr.decode("utf-8", errors="ignore")
        except Exception as e:
            last_err = str(e)

    try:
        os.remove(tmp_path)
    except OSError:
        pass

    if not ok:
        sys.stderr.write(
            "[WARNING] mmdc no pudo renderizar un diagrama Mermaid: "
            f"{last_err or 'Invalid expression'}\n"
        )
    return ok


# ----------------------------------------------------------------------
# Sustituir bloques mermaid por imágenes (o dejar código si fallan)
# ----------------------------------------------------------------------
def replace_mermaid_with_images(
    md_text: str,
    work_dir: Path,
    width: int,
    background: str,
    theme: str,
    fix_subgraphs_flag: bool,
    fix_edges_flag: bool,
    fix_lines_flag: bool,
    escape_parens_flag: bool,
) -> str:
    work_dir.mkdir(parents=True, exist_ok=True)
    counter = 1

    def _apply_fixes(code: str) -> str:
        code = normalize_unicode(code)
        if fix_subgraphs_flag:
            code = fix_mermaid_subgraphs(code)
        if fix_lines_flag:
            code = fix_double_lines(code)
        if escape_parens_flag:
            code = escape_parens_in_labels(code)
        # DESHABILITADAS - rompen sintaxis moderna de Mermaid:
        # if fix_edges_flag:
        #     code = fix_edge_labels(code)
        # code = escape_entities_in_nodes(code)
        return code

    def _repl(match: re.Match) -> str:
        nonlocal counter
        original_code = match.group(1).strip()
        code = _apply_fixes(original_code)

        img_name = f"mermaid_{counter:03d}.png"
        img_path = work_dir / img_name
        idx = counter
        counter += 1

        ok = render_mermaid_to_png(
            code, img_path, width=width, background=background, theme=theme
        )
        if not ok:
            sys.stderr.write(
                f"[WARNING] Mermaid #{idx} inválido o no renderizable. "
                "Se deja como bloque de código en el DOCX.\n"
            )
            return f"```mermaid\n{original_code}\n```"

        abs_path = img_path.resolve()
        # Sin texto alt para evitar pie de nota debajo del diagrama
        return f"![]({abs_path})"

    return MERMAID_BLOCK_RE.sub(_repl, md_text)


def safe_replace_mermaid_with_images(
    md_text: str,
    work_dir: Path,
    width: int,
    background: str,
    theme: str,
    fix_subgraphs_flag: bool,
    fix_edges_flag: bool,
    fix_lines_flag: bool,
    escape_parens_flag: bool,
) -> str:
    """
    Capa de seguridad: si pasa cualquier cosa rara (incluido un 'Invalid expression'
    inesperado en algún lado), NO lanzamos excepción y devolvemos el markdown tal cual.
    """
    try:
        return replace_mermaid_with_images(
            md_text,
            work_dir,
            width,
            background,
            theme,
            fix_subgraphs_flag,
            fix_edges_flag,
            fix_lines_flag,
            escape_parens_flag,
        )
    except Exception as e:
        sys.stderr.write(
            "[WARNING] Error inesperado procesando Mermaid, "
            f"se deja el markdown original sin imágenes. Detalle: {e}\n"
        )
        return md_text


# ----------------------------------------------------------------------
# Estilado de documento
# ----------------------------------------------------------------------
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"


def style_document(doc: Document) -> None:
    styles = doc.styles

    # Estilo Normal
    if "Normal" in styles:
        st = styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(11)

    # Encabezados H1–H3
    for i in range(1, 4):
        name = f"Heading {i}"
        if name in styles:
            h = styles[name]
            h.font.name = "Calibri"
            h.font.size = Pt(14 if i == 1 else 13 if i == 2 else 12)
            h.font.bold = True

    # 1) Corregir párrafos "ID Requerimientos relacionados:"
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if txt.startswith("ID Requerimientos relacionados:"):
            try:
                para.style = styles["Normal"]
            except KeyError:
                pass

    # 2) Forzar negrita en los headings
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.lower().startswith("heading"):
            for run in para.runs:
                if run.bold is None:
                    run.bold = True

    # 3) Redimensionar imágenes grandes
    MAX_W, MAX_H = Cm(15), Cm(23)
    for shape in doc.inline_shapes:
        try:
            w, h = shape.width, shape.height
            if w <= 0 or h <= 0:
                continue
            scale = min(MAX_W / w, MAX_H / h)
            if scale < 1:
                shape.width = int(w * scale)
                shape.height = int(h * scale)
        except Exception:
            pass

    # 4) Aplicar tamaño de fuente 8pt a bloques de código (wireframes) en Sección 12
    in_section_12 = False
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        
        # Detectar inicio de Sección 12
        if "12. Prototipo de Interfaz" in txt or "12. Prototipos de Interfaz" in txt:
            in_section_12 = True
            continue
        
        # Detectar fin de Sección 12 (inicio de otra sección principal)
        if in_section_12 and txt.startswith("13."):
            in_section_12 = False
            continue
        
        # Aplicar tamaño 8pt a bloques de código dentro de Sección 12
        # Los wireframes usan caracteres box-drawing Unicode: ┌─┐│└┘├┤┬┴┼
        if in_section_12:
            # Detectar caracteres box-drawing Unicode (U+2500 a U+257F)
            box_drawing_chars = {'┌', '─', '┐', '│', '└', '┘', '├', '┤', '┬', '┴', '┼', '╔', '═', '╗', '║', '╚', '╝', '╠', '╣', '╦', '╩', '╬'}
            is_wireframe_line = (
                txt.startswith("+") or 
                txt.startswith("|") or 
                txt.startswith("[") or
                "+---" in txt or
                "|---" in txt or
                any(char in txt for char in box_drawing_chars)
            )
            if is_wireframe_line:
                for run in para.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Consolas"  # Fuente monoespaciada para wireframes

    # 5) POST-PROCESAMIENTO: Convertir LINE SEPARATOR (U+2028) en saltos de línea en celdas de tablas
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    LINE_SEP = '\u2028'  # Unicode LINE SEPARATOR
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Procesar cada párrafo de la celda
                for para in cell.paragraphs:
                    full_text = para.text
                    # Verificar si contiene el separador de línea
                    if LINE_SEP in full_text:
                        # Limpiar el párrafo actual
                        for run in para.runs:
                            run.text = ""
                        
                        # Dividir por el separador
                        parts = full_text.split(LINE_SEP)
                        
                        for i, part in enumerate(parts):
                            part = part.strip()
                            
                            if i == 0:
                                # Primera parte: añadir al párrafo existente
                                run = para.add_run(part)
                            else:
                                # Partes siguientes: añadir salto de línea y texto
                                run = para.add_run()
                                br = OxmlElement('w:br')
                                run._r.append(br)
                                run = para.add_run(part)

    # 6) AJUSTE DE TABLAS ESPECÍFICAS
    for table in doc.tables:
        try:
            # Obtenemos datos clave para identificar la tabla
            num_cols = len(table.columns)
            # Texto de la primera celda de la cabecera (en minúsculas para comparar)
            h0 = table.rows[0].cells[0].text.strip().lower()
        except IndexError:
            continue  # Tabla vacía o sin filas

        # --- CASO A: REQUERIMIENTOS FUNCIONALES (3 Columnas: ID, Requerimiento, Descripción) ---
        # Identificador: 3 columnas y la primera se llama "id"
        if num_cols == 3 and h0 == "id":
            table.autofit = False
            table.allow_autofit = False
            
            # Anchos basados en la imagen (ID estrecho, Req medio, Desc ancho)
            w0 = Cm(1.75)  # ID
            w1 = Cm(5.25)  # Requerimiento
            w2 = Cm(9.5)   # Descripción (Resto)

            for row in table.rows:
                row.cells[0].width = w0
                row.cells[1].width = w1
                row.cells[2].width = w2

        # --- CASO B: MATRIZ DE TRAZABILIDAD (3 Columnas: Requisito, HU, Pantalla) ---
        # Identificador: 3 columnas y la primera empieza por "requisito"
        elif num_cols == 3 and h0.startswith("requisito"):
            table.autofit = False
            table.allow_autofit = False
            
            # Anchos basados en la imagen (Requisito ancho, HU medio, Pantalla medio)
            w0 = Cm(8.0)   # Requisito (es la columna con más texto aquí)
            w1 = Cm(4.0)   # Historia Usuario
            w2 = Cm(4.5)   # Pantalla

            for row in table.rows:
                row.cells[0].width = w0
                row.cells[1].width = w1
                row.cells[2].width = w2

        # --- CASO C: INTERFACES DE USUARIO (5 Columnas) ---
        # Identificador: 5 columnas y cabecera empieza por "pantalla"
        elif num_cols == 5 and h0.startswith("pantalla"):
            table.autofit = False
            table.allow_autofit = False
            
            # Anchos basados en la imagen:
            # Pantalla(ID): ~2.5cm | Tipo: ~1.5cm | Desc: ~7cm | Roles: ~2.5cm | HU: ~2.5cm
            w0 = Cm(2.8)  # Pantalla ID
            w1 = Cm(1.7)  # Tipo
            w2 = Cm(7.0)  # Descripción (La más ancha)
            w3 = Cm(2.5)  # Roles
            w4 = Cm(2.5)  # Historia Usuario

            for row in table.rows:
                if len(row.cells) >= 5:
                    row.cells[0].width = w0
                    row.cells[1].width = w1
                    row.cells[2].width = w2
                    row.cells[3].width = w3
                    row.cells[4].width = w4

        # --- CASO D: INTEGRACIONES (4 Columnas: Sistema, Tipo, Propósito, Tecnología) ---
        elif num_cols == 4 and h0 == "sistema":
            table.autofit = False
            table.allow_autofit = False
            
            # Anchos basados en la imagen adjunta
            w0 = Cm(2.8)   # Sistema
            w1 = Cm(2.5)   # Tipo de Integración
            w2 = Cm(8.5)   # Propósito (la más ancha)
            w3 = Cm(2.7)   # Tecnología/Protocolo

            for row in table.rows:
                if len(row.cells) >= 4:
                    row.cells[0].width = w0
                    row.cells[1].width = w1
                    row.cells[2].width = w2
                    row.cells[3].width = w3

        # --- CASO E: PRUEBAS FUNCIONALES (6 Columnas) ---
        elif num_cols == 6 and (h0 == "id prueba" or h0.startswith("id")):
            # Verificar que es la tabla de pruebas (segunda columna = "hu / cu" o similar)
            try:
                h1 = table.rows[0].cells[1].text.strip().lower()
            except IndexError:
                h1 = ""
            
            if "hu" in h1 or "cu" in h1:
                table.autofit = False
                table.allow_autofit = False
                
                # Anchos exactos especificados por el usuario
                w0 = Cm(1.69)  # ID Prueba
                w1 = Cm(2.0)   # HU / CU
                w2 = Cm(2.22)  # Pantalla(s)
                w3 = Cm(2.28)  # Actor
                w4 = Cm(2.75)  # Criterios de Aceptación
                w5 = Cm(5.03)  # Resultado Esperado

                for row in table.rows:
                    if len(row.cells) >= 6:
                        row.cells[0].width = w0
                        row.cells[1].width = w1
                        row.cells[2].width = w2
                        row.cells[3].width = w3
                        row.cells[4].width = w4
                        row.cells[5].width = w5

        # --- CASO F: MATRIZ DE TRAZABILIDAD MODO COMPLETO (5 Columnas) ---
        elif num_cols == 5 and h0.startswith("requisito"):
            table.autofit = False
            table.allow_autofit = False
            
            # Anchos basados en la imagen adjunta para modo completo
            w0 = Cm(5.0)   # Requisito
            w1 = Cm(2.2)   # Historia(s) de Usuario
            w2 = Cm(2.0)   # Caso(s) de Uso
            w3 = Cm(2.3)   # Pantalla(s)
            w4 = Cm(5.0)   # Prueba(s)

            for row in table.rows:
                if len(row.cells) >= 5:
                    row.cells[0].width = w0
                    row.cells[1].width = w1
                    row.cells[2].width = w2
                    row.cells[3].width = w3
                    row.cells[4].width = w4

# ----------------------------------------------------------------------
# Conversión markdown -> docx
# ----------------------------------------------------------------------
def convert_md_to_docx(
    md_path: Path,
    out_path: Path,
    width: int,
    background: str,
    theme: str,
    fix_subgraphs_flag: bool,
    fix_edges_flag: bool,
    fix_lines_flag: bool,
    escape_parens_flag: bool,
) -> None:
    original = md_path.read_text(encoding="utf-8")
    original = clean_markdown_text(original)
    original = fix_acceptance_criteria_formatting(original)
    with tempfile.TemporaryDirectory() as td:
        tdir = Path(td)

        # USAMOS LA VERSIÓN "SAFE": nunca lanza, aunque haya 'Invalid expression'
        md_with_imgs = safe_replace_mermaid_with_images(
            original,
            tdir,
            width,
            background,
            theme,
            fix_subgraphs_flag,
            fix_edges_flag,
            fix_lines_flag,
            escape_parens_flag,
        )

        tmp_docx = tdir / "tmp.docx"
        pypandoc.convert_text(
            md_with_imgs,
            "docx",
            format="markdown-yaml_metadata_block",
            outputfile=str(tmp_docx),
            extra_args=["--standalone"],
        )
        doc = Document(str(tmp_docx))
        style_document(doc)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))


# ----------------------------------------------------------------------
# CLI
# ----------------------------------------------------------------------
def parse_args(argv=None):
    ap = argparse.ArgumentParser(
        description="Convierte Markdown a DOCX renderizando Mermaid con normalizaciones robustas (v2.3)."
    )
    ap.add_argument(
        "input_md",
        help="Fichero .md o carpeta con varios .md (01_xxx.md, 02_xxx.md, ...)",
    )
    ap.add_argument("output_docx", help="Fichero DOCX de salida")
    ap.add_argument("--width", type=int, default=1600)
    ap.add_argument("--background", default="white")
    ap.add_argument("--theme", default="default")
    ap.add_argument("--no-fix-subgraphs", dest="fix_subgraphs", action="store_false")
    ap.add_argument("--no-fix-edges", dest="fix_edges", action="store_false")
    ap.add_argument("--no-fix-lines", dest="fix_lines", action="store_false")
    ap.add_argument(
        "--no-escape-parens", dest="escape_parens", action="store_false"
    )
    ap.set_defaults(
        fix_subgraphs=True,
        fix_edges=True,
        fix_lines=True,
        escape_parens=True,
    )
    return ap.parse_args(argv)


def main(argv=None):
    args = parse_args(argv)
    in_path = Path(args.input_md)
    out = Path(args.output_docx)
    if not in_path.exists():
        sys.exit(f"[ERROR] No existe el archivo o carpeta: {in_path}")

    # OJO: aquí ya NO atrapamos excepciones para no reescribir el mensaje ('Invalid expression')
    if in_path.is_dir():
        combined_md = build_combined_markdown_from_dir(in_path)
        with tempfile.TemporaryDirectory() as td:
            tmp_md = Path(td) / "combined.md"
            tmp_md.write_text(combined_md, encoding="utf-8")
            convert_md_to_docx(
                tmp_md,
                out,
                args.width,
                args.background,
                args.theme,
                args.fix_subgraphs,
                args.fix_edges,
                args.fix_lines,
                args.escape_parens,
            )
    else:
        convert_md_to_docx(
            in_path,
            out,
            args.width,
            args.background,
            args.theme,
            args.fix_subgraphs,
            args.fix_edges,
            args.fix_lines,
            args.escape_parens,
        )

    print(f"[OK] DOCX generado en: {out}")


if __name__ == "__main__":
    main()
